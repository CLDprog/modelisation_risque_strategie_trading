"""
Génère l'AUDIT DES MANQUES (gap analysis) — Roadmap v4.0 vs projet.
Document Word exhaustif, factuel, avec preuves (fichier:ligne) et gravité.

Base corrigée : le dossier docs/ EXISTE (étape 16 réévaluée Conforme-).
Audit indépendant — état au 2026-06-05.

Usage : python scripts/generate_gap_audit.py
Sortie : docs/Audit_conformite_roadmap_v4_detaille.docx
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Charte couleurs (reprise de generate_audit_v2.py)
# ----------------------------------------------------------------------------
C_CONFORME = "2E7D32"   # vert
C_PARTIEL  = "E65100"   # orange
C_MINEUR   = "B7950B"   # ambre
C_MANQUANT = "B71C1C"   # rouge
C_HEADER   = "1F3864"   # bleu foncé
C_ACCENT   = "2E75B6"   # bleu accent
C_LIGHT    = "EAF0F7"   # fond clair

STATUS_COLOR = {
    "Conforme": C_CONFORME, "Conforme-": C_CONFORME,
    "Partiel": C_PARTIEL, "Faible": C_PARTIEL, "Manquant": C_MANQUANT,
}
GRAV_COLOR = {"Manquant": C_MANQUANT, "Partiel": C_PARTIEL, "Mineur": C_MINEUR}


# ----------------------------------------------------------------------------
# Helpers bas niveau
# ----------------------------------------------------------------------------
def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=9, align="left",
                  white=False, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Consolas" if mono else "Calibri"
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_fixed_layout(t):
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)


def vcell(cell):
    """Centre verticalement le contenu d'une cellule."""
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(C_HEADER)
        run.font.name = "Calibri"
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(C_ACCENT)
        run.font.name = "Calibri"
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(C_HEADER)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def para(doc, text, bold=False, italic=False, size=10.5, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(doc, text, size=10, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor.from_string("444444")
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = "[Clic droit dans la table → « Mettre à jour les champs » pour générer le sommaire]"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, t, fld_end):
        run._r.append(el)


def add_page_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pre = p.add_run("Audit de conformité — Vol Risk Infrastructure   ·   Page ")
    pre.font.size = Pt(8); pre.font.color.rgb = RGBColor.from_string("888888")
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r = p.add_run(); r._r.append(fld1); r._r.append(instr); r._r.append(fld2)
    r.font.size = Pt(8)


def status_table(doc, rows, col_widths, headers):
    t = doc.add_table(rows=0, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    set_fixed_layout(t)
    hdr = t.add_row().cells
    for i, label in enumerate(headers):
        set_cell_text(hdr[i], label, bold=True, white=True, size=9, align="center")
        shade(hdr[i], C_HEADER); vcell(hdr[i])
    for row in rows:
        c = t.add_row().cells
        for i, val in enumerate(row):
            if i == len(headers) - 1:  # dernière colonne = statut
                set_cell_text(c[i], val, bold=True, white=True, size=8.5, align="center")
                shade(c[i], STATUS_COLOR.get(val, C_ACCENT))
            else:
                set_cell_text(c[i], val, size=8.5)
            vcell(c[i])
    for row in t.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)
    return t


def gap_table(doc, headers, rows, col_widths):
    t = doc.add_table(rows=0, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    set_fixed_layout(t)
    hdr = t.add_row().cells
    for i, label in enumerate(headers):
        set_cell_text(hdr[i], label, bold=True, white=True, size=8.5, align="center")
        shade(hdr[i], C_HEADER); vcell(hdr[i])
    for n, row in enumerate(rows):
        c = t.add_row().cells
        for i, val in enumerate(row):
            is_grav = (i == len(headers) - 1)
            is_proof = (headers[i].startswith("Preuve"))
            if is_grav:
                set_cell_text(c[i], val, bold=True, white=True, size=8, align="center")
                shade(c[i], GRAV_COLOR.get(val, C_ACCENT))
            else:
                set_cell_text(c[i], val, size=8, mono=is_proof)
                if n % 2 == 1:
                    shade(c[i], C_LIGHT)
            vcell(c[i])
    for row in t.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)
    return t


# ============================================================================
# CONSTRUCTION DU DOCUMENT
# ============================================================================
doc = Document()

# Style par défaut
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# Page A4, marges 0.7"
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)
add_page_footer(doc)

CW = 6.9  # largeur de contenu approximative (pouces)

# ----------------------------------------------------------------------------
# PAGE DE TITRE
# ----------------------------------------------------------------------------
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AUDIT DE CONFORMITÉ"); r.bold = True; r.font.size = Pt(30)
r.font.color.rgb = RGBColor.from_string(C_HEADER); r.font.name = "Calibri"

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Analyse exhaustive des manques (gap analysis)")
r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(C_ACCENT); r.font.name = "Calibri"

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Vol Risk Infrastructure  vs  Industrial Roadmap for a Volatility Infrastructure Platform (v4.0)")
r.font.size = Pt(11); r.italic = True; r.font.name = "Calibri"

# trait
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), C_ACCENT)
pbdr.append(bottom); pPr.append(pbdr)

for _ in range(2):
    doc.add_paragraph()

meta = doc.add_table(rows=0, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for label, val in [
    ("Périmètre", "16 étapes + framework mathématique + annexes (51 pages de roadmap)"),
    ("Base de code auditée", "architecture_risque/ — ~30 modules, 56 tests, configs, docs/"),
    ("Méthode", "Lecture du code + grep sur l'intégralité du repo (src, pages, configs, tests, docs)"),
    ("Verdict global", "≈ 9 Conforme · 7 Partiel · 0 Faible (sur 16 étapes)"),
    ("Date", date.today().isoformat()),
]:
    c = meta.add_row().cells
    set_cell_text(c[0], label, bold=True, size=9.5, color=C_HEADER, align="right")
    set_cell_text(c[1], val, size=9.5)
    c[0].width = Inches(1.8); c[1].width = Inches(5.0)

doc.add_paragraph()
para(doc, "Note de cadrage. Ce document liste ce qui MANQUE par rapport à la roadmap. "
          "Les briques déjà conformes (framework mathématique Eq.1–21/24/25, solveur d'IV, "
          "pricing BS-76 + CRR, architecture 2-process, lineage, docs opérationnelles) ne sont "
          "rappelées que brièvement. Chaque manque est étayé par une preuve (fichier:ligne) et "
          "classé : Manquant (rien) · Partiel (existe mais incomplet) · Mineur.",
     italic=True, size=9.5, space_after=4)
para(doc, "Correction de méthode. Une première passe avait conclu à tort à l'absence du dossier "
          "docs/ (erreur de glob). Après vérification, docs/ existe et couvre l'essentiel de "
          "l'étape 16 ; le présent rapport intègre cette correction.",
     italic=True, size=9.5, color=C_PARTIEL)

doc.add_page_break()

# ----------------------------------------------------------------------------
# SOMMAIRE
# ----------------------------------------------------------------------------
h1(doc, "Sommaire")
add_toc(doc)
doc.add_page_break()

# ----------------------------------------------------------------------------
# 1. SYNTHÈSE EXÉCUTIVE
# ----------------------------------------------------------------------------
h1(doc, "1. Synthèse exécutive")

para(doc, "Le socle quantitatif et l'architecture sont solides — souvent au-dessus du minimum "
          "exigé (no-arbitrage papillon, inversion d'IV américaine, anomaly detection, isolation "
          "collecteur/dashboard). L'écart de conformité se concentre sur trois familles :", space_after=6)
bullet(doc, "la PERSISTANCE : 3 tables du modèle de données ne sont jamais écrites, et plusieurs "
            "diagnostics (candidats de forward, points rejetés, paramètres SVI) sont calculés mais perdus.",
       bold_prefix="Persistance — ")
bullet(doc, "les RÉCONCILIATIONS : ni finite-difference vs analytique, ni Greeks broker vs plateforme, "
            "alors que les deux sont exigées explicitement (étape 11 + QC « Greek sanity »).",
       bold_prefix="Réconciliations — ")
bullet(doc, "l'INDUSTRIALISATION : pas de scheduler système, pas de correlation IDs, pas de routage "
            "d'alertes externe, et seule la couche 1 (tests unitaires) des 4 couches de tests existe.",
       bold_prefix="Industrialisation — ")

para(doc, "Décompte des manques", bold=True, size=11.5, space_after=2)
dec = status_table(
    doc,
    [["Manquant (aucune implémentation)", "16", "Manquant"],
     ["Partiel (existe mais incomplet)", "≈ 10", "Partiel"],
     ["Mineur", "≈ 12", "Mineur"]],
    [4.0, 1.5, 1.4],
    ["Catégorie de manque", "Nombre", "Gravité"],
)

para(doc, "")
para(doc, "Comparaison avec l'auto-audit du projet. Vos scripts internes "
          "(scripts/generate_audit_v2.py) classent 14 étapes « Conforme » et 2 « Partiel ». "
          "L'audit indépendant est plus sévère (≈ 9 Conforme / 7 Partiel) : plusieurs étapes "
          "marquées « Conforme » comportent des trous vérifiables — notamment l'étape 9 "
          "(« grille + params » alors que les paramètres SVI ne sont jamais persistés) et l'étape 11 "
          "(« réconciliation broker » revendiquée alors qu'elle n'existe pas dans le code).",
     size=10, space_after=6)

doc.add_page_break()

# ----------------------------------------------------------------------------
# 2. CE QUI EST COMPLET (ACQUIS)
# ----------------------------------------------------------------------------
h1(doc, "2. Ce qui est complet (acquis du projet)")
para(doc, "Vue synthétique de ce qui est livré et opérationnel — la majorité de la pile. "
          "À lire avant la liste des manques : l'ossature data → forward → IV → surface → "
          "pricing → risk → scénarios → QC est en place de bout en bout.", space_after=8)

h2(doc, "2.1  Framework mathématique — 23 des 25 équations implémentées")
status_table(
    doc,
    [["Spot, forward & carry", "Eq.1 – Eq.5", "Conforme"],
     ["Log-moneyness & variance totale", "Eq.6 – Eq.7", "Conforme"],
     ["Pricing européen (Black-76)", "Eq.8 – Eq.11", "Conforme"],
     ["Pricing américain (arbre CRR)", "Eq.12", "Conforme"],
     ["Greeks analytiques (Δ, Γ, ν, Θ, $Γ, $ν)", "Eq.13 – Eq.18", "Conforme"],
     ["Approximation de PnL par les Greeks", "Eq.19", "Conforme"],
     ["Surface SVI & monotonie calendaire", "Eq.20 – Eq.21", "Conforme"],
     ["Stats robustes (MAD) & diagnostics spread/mid", "Eq.24 – Eq.25", "Conforme"]],
    [3.5, 2.1, 1.0],
    ["Bloc mathématique", "Équations", "Statut"],
)
para(doc, "Seules manquent Eq.22 (interpolation de variance cross-maturité) et Eq.23 "
          "(variance de panier) — soit 23/25.", italic=True, size=9, space_after=8)

h2(doc, "2.2  Couches & étapes opérationnelles")
status_table(
    doc,
    [["Connectivité (1)", "Machine d'état 5 états, backoff exponentiel + jitter, heartbeat", "Conforme"],
     ["Instrument master (2)", "Clés canoniques, round-trip from_key, discovery IBKR, SQLite", "Conforme"],
     ["Ingestion (3)", "Couche brute immuable append-only, validation, reconnexion", "Conforme"],
     ["Snapshots (5)", "Builder pur déterministe, fallback labellisé, flags stale", "Conforme"],
     ["Forward (6)", "Parité put-call, pondération liquidité, rejet MAD, score de confiance", "Conforme"],
     ["Solveur d'IV (8)", "Brent bracketé européen + CRR américain, diagnostics complets, batch", "Conforme"],
     ["Surface (9)", "SVI par tranche + fallback spline + monotonie + no-arb papillon", "Conforme"],
     ["Pricing (10)", "Black-76 + Greeks analytiques + arbre CRR, objets typés, conventions", "Conforme"],
     ["Risk (11)", "Greeks ligne-à-ligne, positions IBKR réelles, monétisation $Γ/$ν", "Conforme"],
     ["Scénarios (12)", "Grille versionnée (7), repricing complet + approx Greeks, contributeurs", "Conforme"],
     ["Replay (13)", "Même code path que le live, reconstruction depuis le raw, compare replay/live", "Conforme"],
     ["QC & validation (14)", "Checks nommés, baselines glissantes, anomaly detection, table de triage", "Conforme"],
     ["Observabilité (15)", "Isolation 2-process, alerts.json, statut collecteur, dashboard 10 pages", "Conforme"],
     ["Documentation (16)", "5 runbooks, release checklist, known limitations, interface contracts", "Conforme"]],
    [1.9, 4.0, 1.0],
    ["Couche / étape", "Ce qui est livré et opérationnel", "Statut"],
)

h2(doc, "2.3  Au-dessus du minimum exigé (différenciateurs à valoriser)")
for txt in [
    "No-arbitrage papillon (convexité des prix de calls en strike) — non exigé au minimum roadmap.",
    "Inversion d'IV américaine via arbre CRR — au-delà du simple proxy demandé pour les options à exercice anticipé.",
    "Détection d'anomalies + baselines glissantes (z-score MAD) + table de triage — au-delà des checks ponctuels.",
    "Architecture 2-process (collecteur / dashboard) avec écritures atomiques — « process isolation » de la roadmap.",
    "Lineage complet (code_version, config_hash, run_id) attaché à TOUTES les sorties dérivées.",
    "Données IBKR réelles (différé gratuit) + repli automatique sur le close historique marché fermé.",
    "56 tests unitaires : pricing (parité, intrinsèque, limites), round-trip IV européen + américain, forward, no-arb papillon, anomalies, replay vs live.",
]:
    bullet(doc, txt, size=10)

doc.add_page_break()

# ----------------------------------------------------------------------------
# 3. TABLEAU DE SYNTHÈSE DES 16 ÉTAPES
# ----------------------------------------------------------------------------
h1(doc, "3. Tableau de synthèse — 16 étapes")
para(doc, "Statut indépendant et écart principal par étape. Le détail figure en section 4.",
     italic=True, size=9.5, space_after=6)

steps_rows = [
    ["1 — Accès, env., sécurité", "deps Dash absentes de requirements.txt ; pas de check clock-sync/entitlement", "Conforme-"],
    ["2 — Instrument master", "payloads broker bruts non persistés (« evidence » exigée)", "Conforme-"],
    ["3 — Ingestion market data", "exchange_ts toujours None ; events malformés non mis en quarantaine", "Conforme"],
    ["4 — Stockage & data model", "3 tables jamais écrites ; aucune politique de rétention", "Partiel"],
    ["5 — Spot builder / snapshots", "fallback bid au lieu de close/carry-forward du snapshot précédent", "Conforme"],
    ["6 — Forward & carry", "diagnostics par strike non persistés ; pas de lissage terme ni courbe de taux", "Partiel"],
    ["7 — Normalisation / QC quotes", "min_open_interest jamais appliqué ; checks non « nommés »", "Partiel"],
    ["8 — Solveur d'IV", "(fort) Newton annoncé mais non codé ; coord. delta absente", "Conforme"],
    ["9 — Surface engine", "surface_parameters jamais écrits ; Eq.22 (interp. cross-maturité) absente", "Partiel"],
    ["10 — Pricing engine", "pas d'API vectorisée ; table pricing_results absente", "Conforme"],
    ["11 — Greeks & risk", "aggregate_risk() jamais appelé ; réconciliations broker + finite-diff absentes", "Partiel"],
    ["12 — Scénarios", "complet (rotations de courbe optionnelles non couvertes)", "Conforme"],
    ["13 — Reconstruction / replay", "réécrit la même partition au lieu de versionner physiquement", "Conforme-"],
    ["14 — Validation framework", "3 checks nommés manquants ; escalade documentée mais non codée", "Partiel"],
    ["15 — Orchestration / observabilité", "pas de scheduler système ; pas de correlation IDs ; alertes locales", "Partiel"],
    ["16 — Handover production", "docs/ solide ; manquent module READMEs + diagramme + schemas.md détaillé", "Conforme-"],
]
status_table(doc, steps_rows, [2.3, 3.6, 1.0],
             ["Étape", "Écart principal", "Statut"])

doc.add_page_break()

# ----------------------------------------------------------------------------
# 3. DÉTAIL PAR ÉTAPE
# ----------------------------------------------------------------------------
h1(doc, "4. Détail des manques, étape par étape")

def step_detail(num_title, statut, realise, manques):
    h2(doc, num_title)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    rl = p.add_run("Statut : "); rl.bold = True; rl.font.size = Pt(10.5); rl.font.name = "Calibri"
    rs = p.add_run(statut); rs.bold = True; rs.font.size = Pt(10.5); rs.font.name = "Calibri"
    rs.font.color.rgb = RGBColor.from_string(STATUS_COLOR.get(statut, C_ACCENT))
    pr = doc.add_paragraph(); pr.paragraph_format.space_after = Pt(3)
    rl = pr.add_run("Déjà en place. "); rl.bold = True; rl.font.size = Pt(9.5)
    rl.font.color.rgb = RGBColor.from_string(C_CONFORME); rl.font.name = "Calibri"
    rt = pr.add_run(realise); rt.font.size = Pt(9.5); rt.font.name = "Calibri"
    pm = doc.add_paragraph(); pm.paragraph_format.space_after = Pt(2)
    rl = pm.add_run("Ce qui manque :"); rl.bold = True; rl.font.size = Pt(9.5)
    rl.font.color.rgb = RGBColor.from_string(C_MANQUANT); rl.font.name = "Calibri"
    for grav, txt, proof in manques:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        tag = p.add_run(f"[{grav}] "); tag.bold = True; tag.font.size = Pt(9)
        tag.font.color.rgb = RGBColor.from_string(GRAV_COLOR.get(grav, C_ACCENT)); tag.font.name = "Calibri"
        rt = p.add_run(txt); rt.font.size = Pt(9.5); rt.font.name = "Calibri"
        if proof:
            rp = p.add_run(f"  ({proof})"); rp.font.size = Pt(8.5); rp.font.name = "Consolas"
            rp.font.color.rgb = RGBColor.from_string("666666")


step_detail(
    "Étape 1 — Accès, environnements et sécurité", "Conforme-",
    "Smoke test bootstrap.py (résout SPY, lit une quote, écrit un event), config externalisée "
    "(YAML), secrets hors dépôt (.env), IBKRSession avec backoff exponentiel + jitter, "
    "versions épinglées.",
    [("Partiel", "dash et dash-bootstrap-components absents de requirements.txt alors que "
                 "l'app front en dépend — install KO sur machine neuve. docs/environment.md "
                 "affirme pourtant qu'ils sont installés (incohérence doc/code).",
      "requirements.txt vs docs/environment.md:18"),
     ("Mineur", "Health checks incomplets : pas de vérification de synchronisation d'horloge "
                 "ni d'entitlement market-data dans le bootstrap (exigés étape 1).",
      "scripts/bootstrap.py"),
     ("Mineur", "Pas de lock file (poetry.lock / pip-tools) ; reproductibilité reposant sur "
                 "des versions épinglées seulement.", "requirements.txt")],
)

step_detail(
    "Étape 2 — Instrument master & discovery", "Conforme-",
    "Dataclasses Underlying/OptionContract, instrument_key canonique, from_key (round-trip), "
    "InstrumentMaster, discovery IBKR, persistance SQLite, validate() (multiplicateurs/strikes), "
    "déduplication.",
    [("Manquant", "Payloads broker bruts non persistés. La roadmap insiste (« store raw contract "
                  "payloads as evidence ») : discover_universe_from_ibkr crée les contrats mais "
                  "ne sauvegarde jamais la réponse brute reqSecDefOptParams.",
      "src/universe/contracts.py:225")],
)

step_detail(
    "Étape 3 — Ingestion market data", "Conforme",
    "RawMarketEvent canonique, RawEventWriter append-only (batch/flush), validation des champs, "
    "souscription underlyings+options, reconnexion/heartbeat, coverage_summary, couche brute "
    "immuable alimentée avant tout calcul.",
    [("Mineur", "exchange_ts présent dans le schéma mais toujours None — jamais renseigné dans "
                "le chemin live (seul receipt_ts l'est).", "src/collectors/raw_writer.py:40,55"),
     ("Mineur", "Events malformés comptés (_malformed_count) puis droppés, pas mis en "
                "quarantaine avec reason code dans un store (« quarantine with a reason code »).",
      "src/collectors/raw_writer.py:102"),
     ("Mineur", "« Missing intervals » non calculés dans les résumés de session.",
      "src/collectors/raw_writer.py:225")],
)

step_detail(
    "Étape 4 — Stockage persistant & data model", "Partiel",
    "ParquetStore partitionné par date, MetadataStore SQLite (instrument_master, job_manifests, "
    "qc_results), lineage (code_version/config_hash/run_id) sur toutes les sorties dérivées, "
    "manifests JSON + DB, écritures atomiques.",
    [("Manquant", "Table surface_parameters jamais écrite (paramètres SVI a,b,ρ,m,σ par tranche).",
      "0 occurrence de store.write(\"surface_parameters\")"),
     ("Manquant", "Table pricing_results jamais écrite (prix + Greeks par contrat).",
      "0 occurrence"),
     ("Manquant", "Table positions jamais écrite (positions source-of-record versionnées).",
      "0 occurrence"),
     ("Mineur", "Aucune politique de rétention (Tier 1–4 de la Part XV) ; partition par date "
                "seule (la roadmap suggère date × underlying × layer).", "src/storage/schemas.py"),
     ("Mineur", "Règles d'évolution de schéma / write-ahead validation non formalisées.", "")],
)

step_detail(
    "Étape 5 — Spot builder & snapshots", "Conforme",
    "build_snapshot pur et déterministe, choose_reference_spot (mid→last→fallback), reference_type "
    "stocké, flags stale/fallback, métriques de complétude, rejouable depuis le raw.",
    [("Mineur", "Ordre de fallback partiel : le builder fait mid→last→bid_fallback ; la roadmap "
                "demande mid→last→close→carry-forward du dernier snapshot fiable. Le report du "
                "snapshot précédent (carry-forward) n'est pas implémenté.",
      "src/snapshots/builder.py:78"),
     ("Mineur", "En live, reference_type est figé à « live_delayed » — la granularité mid/last/close "
                "exigée par le data dictionary est perdue.", "src/data/live.py / run_collector.py:388")],
)

step_detail(
    "Étape 6 — Forward & implied carry", "Partiel",
    "parity_forward (Eq.2), liquidity_weight, weighted_mean (Eq.4), implied_carry (Eq.5), "
    "robust_zscore_mad (Eq.24), score de confiance, seuil min de candidats, quality flags.",
    [("Partiel", "Diagnostics par candidat NON persistés : forward_result_to_dict supprime la liste "
                 "des ForwardCandidate (strike, résidu de parité, poids). La roadmap exige « keep "
                 "every intermediate candidate row for audit ».", "src/forwards/engine.py:240"),
     ("Partiel", "Pas de lissage de la structure par terme des forwards ; chaque maturité est "
                 "estimée indépendamment.", "src/forwards/engine.py"),
     ("Partiel", "Taux constant (0.053) — pas de courbe de taux ; le carry implicite n'est comparé "
                 "à aucune attente (diagnostic absent).", "configs/pricing.yaml"),
     ("Mineur", "Politique de fallback implicite (revert vers tous les candidats / médiane) plutôt "
                "qu'un objet fallback policy explicite.", "src/forwards/engine.py:201")],
)

step_detail(
    "Étape 7 — Normalisation & QC des quotes", "Partiel",
    "Flag is_usable + reject_reason par quote, détection spread anormal / crossed-locked / stale, "
    "conservation de la chaîne complète ET filtrée (rejets auditables).",
    [("Partiel", "min_open_interest (=10 dans qc.yaml) JAMAIS appliqué dans le chemin live — aucune "
                 "quote n'est rejetée sur l'open interest.", "src/data/live.py:341"),
     ("Partiel", "Les décisions de rejet sont des if/elif inline, pas une librairie de checks "
                 "nommés et journalisés séparément (« do not implement QC as a monolithic "
                 "if-statement »).", "src/data/live.py:347"),
     ("Mineur", "Pas de rejet d'outliers sur les IV préliminaires au stade quote (l'outlier "
                "rejection n'existe que dans le forward engine).", "")],
)

step_detail(
    "Étape 8 — Solveur d'IV", "Conforme",
    "Brent bracketé (européen) + inversion AMÉRICAINE via CRR, bornes intrinsèque/supérieure, "
    "IvSolveResult complet (converged, iterations, residual, bounds, failure_reason, model), "
    "batch solve_chain_iv, log-moneyness + variance totale. Round-trip testé.",
    [("Mineur", "Accélérateur Newton « inside a safe bracket » annoncé dans le docstring mais non "
                "implémenté (Brent pur).", "src/iv/solver.py:5"),
     ("Mineur", "Coordonnée delta (« delta if available ») absente de la table iv_points.",
      "src/iv/solver.py:285")],
)

step_detail(
    "Étape 9 — Surface engine", "Partiel",
    "SVI par tranche (Eq.20), fallback spline, monotonie calendaire (Eq.21), no-arbitrage papillon "
    "(au-dessus du minimum), RMSE/max_error, grille reconstruite, robustesse DataFrame vide.",
    [("Manquant", "surface_parameters (a,b,ρ,m,σ) calculés mais JAMAIS persistés : "
                  "surface_to_dataframe n'écrit que la grille. La roadmap exige « save both fitted "
                  "parameters AND reconstructed grid values ».",
      "calc src/surfaces/calibration.py:127 ; écriture absente"),
     ("Manquant", "Eq.22 — interpolation de variance cross-maturité absente : aucune fonction pour "
                  "la vol à une maturité arbitraire entre deux tranches (surface slice-par-slice "
                  "seulement).", "0 occurrence d'interp. cross-maturité"),
     ("Partiel", "Points IV rejetés et warnings de fit non stockés (« store rejected points and "
                 "fit warnings »).", "src/surfaces/calibration.py:288")],
)

step_detail(
    "Étape 10 — Pricing engine", "Conforme",
    "Black-76 européen + Greeks analytiques (Eq.8–18), arbre CRR américain, objets résultats typés, "
    "tests de référence (convergence américain→européen, parité), conventions documentées.",
    [("Mineur", "Pas d'API de pricing vectorisée (« scalar AND vectorized interfaces ») — calcul "
                "scalaire uniquement.", "src/pricing/european.py"),
     ("Manquant", "Table pricing_results non matérialisée (prix+Greeks par contrat) — voir étape 4.", ""),
     ("Mineur", "Pas de tests de performance / benchmark fixtures.", "tests/")],
)

step_detail(
    "Étape 11 — Greeks & per-position risk", "Partiel",
    "compute_position_risk (prix+Greeks+monétisés ligne à ligne), Position, enrich_portfolio_greeks "
    "depuis le portefeuille IBKR, IV résolue depuis le prix de marché, dollar gamma/vega.",
    [("Partiel", "aggregate_risk() défini mais JAMAIS appelé (dead code) — aucun agrégat groupé "
                 "produit ni persisté. La page Greeks fait un .sum() à plat, sans groupby "
                 "underlying/maturité/desk. La table « risk_aggregates » contient en réalité du "
                 "ligne-à-ligne.", "def aggregation.py:133 ; somme pages/greeks.py:121"),
     ("Manquant", "Réconciliation Greeks broker vs plateforme absente : les Greeks broker sont "
                  "captés dans les raw events mais jamais comparés (« reconcile against "
                  "broker-returned Greeks if available »).",
      "captés raw_writer.py:209 ; 0 réconciliation"),
     ("Manquant", "Réconciliation finite-difference vs analytique absente, alors qu'exigée "
                  "(« reconciles against finite-difference checks ») : les bumps de pricing.yaml "
                  "ne sont utilisés nulle part.", "configs/pricing.yaml:12 ; 0 usage")],
)

step_detail(
    "Étape 12 — Scénario engine", "Conforme",
    "Scenario typé, grille versionnée (7 scénarios dans scenarios.yaml), repricing complet "
    "(source de vérité) + approximation Greeks (Eq.19), top contributeurs, attribution par ligne, "
    "scenario_results persistés.",
    [("Mineur", "Familles optionnelles non couvertes (rotations de courbe / skew twists) — "
                "explicitement « layerable later » par la roadmap, donc acceptable à ce stade.", "")],
)

step_detail(
    "Étape 13 — Reconstruction historique & replay", "Conforme-",
    "replay_pipeline sur plage de dates (même code path que le live), reconstruction autonome de "
    "l'univers depuis les clés brutes, gestion des partitions manquantes, compare_replay_vs_live, "
    "manifest par run.",
    [("Partiel", "Versionnement physique manquant : replay réécrit le MÊME data.parquet du dt "
                 "(write écrase). La roadmap exige « write to versioned historical partitions "
                 "rather than overwriting ». Le run_id est dans les lignes mais le fichier est "
                 "écrasé.", "src/orchestration/jobs.py:432 / schemas.py:40"),
     ("Mineur", "compare_replay_vs_live existe mais n'est pas branché automatiquement au job de "
                "replay.", "src/orchestration/jobs.py:466")],
)

step_detail(
    "Étape 14 — Validation framework & anomalies", "Partiel",
    "Checks nommés (QcResult status/severity/measured/threshold/reason_code/context), baselines "
    "MAD, détection d'anomalies (z robuste), table de triage, historique QC — au-dessus du minimum.",
    [("Manquant", "Check « option chain coverage » (nb min de calls/puts éligibles par maturité) "
                  "absent.", "src/qc/checks.py"),
     ("Manquant", "Check « parity residual » (résidu de parité par strike) absent en tant que "
                  "check QC nommé.", "src/qc/checks.py"),
     ("Manquant", "Check « Greek sanity » (finite-diff vs analytique) absent — voir étape 11.", ""),
     ("Partiel", "Politique d'escalade S1–S4 documentée (docs/runbooks.md) mais non codée comme "
                 "routage (mapping sévérité → owner → SLA).", "docs/runbooks.md:111")],
)

step_detail(
    "Étape 15 — Orchestration, logging & observabilité", "Partiel",
    "Boucle collecteur always-on, isolation collecteur/dashboard, manifests, alerts.json, "
    "collector_status.json, logs structurés (loguru), reconnexion, dashboard Dash 10 pages "
    "(3 couches de santé).",
    [("Manquant", "Pas de scheduler système (cron/systemd/Airflow) — boucle in-process seulement. "
                  "Limitation assumée dans docs/known_limitations.md.", "run_collector.py:512"),
     ("Manquant", "Correlation IDs absents : collector_session_id et run_id existent mais ne sont "
                  "pas tissés ensemble dans les logs pour relier sessions de collecte et jobs.", ""),
     ("Manquant", "Routage d'alertes externe (email/Slack) absent — alerts.json local seulement "
                  "(limitation assumée).", "run_collector.py:393"),
     ("Partiel", "Catalogue de métriques (Part XIV) partiel : quelques métriques dans le statut, "
                 "mais pas l'ensemble (event rates, stale ratios, bounds-hit, runtimes…).", "")],
)

step_detail(
    "Étape 16 — Handover production", "Conforme-",
    "docs/ couvre l'essentiel : environment.md, runbooks.md (5 runbooks SOD/intraday/EOD/replay/"
    "incident + sévérités), release_checklist.md (catégories A/B/C + rollback), known_limitations.md, "
    "interface_contracts.md (signatures gelées + schémas).",
    [("Mineur", "Pas de README par module (« every module should have a README »).", "docs/"),
     ("Mineur", "Pas de diagramme d'architecture dédié (architecture_overview) — seulement de "
                "l'ASCII dans README/CLAUDE.", "docs/"),
     ("Mineur", "schemas.md pas en version champ-par-champ + règles de partitionnement (le cœur "
                "est dans interface_contracts.md).", "docs/interface_contracts.md")],
)

doc.add_page_break()

# ----------------------------------------------------------------------------
# 4. LISTE EXHAUSTIVE PAR THÈME
# ----------------------------------------------------------------------------
h1(doc, "5. Liste exhaustive des manques, par thème")
para(doc, "Vue transverse (recoupe la section 3). Chaque ligne porte sa preuve et sa gravité.",
     italic=True, size=9.5, space_after=8)

GH = ["Élément manquant", "Exigé par (roadmap)", "Preuve dans le code", "Gravité"]
GW = [1.9, 1.5, 2.6, 0.9]

h2(doc, "A. Briques mathématiques")
gap_table(doc, GH, [
    ["Eq.22 — interpolation de variance cross-maturité", "Part II + pseudo-code surface",
     "interp. seulement intra-tranche (spline) — calibration.py:19", "Manquant"],
    ["Eq.23 — identité de variance de panier / indice", "Part II (section dédiée)",
     "« basket » = 0 occurrence dans tout le repo", "Manquant"],
], GW)

h2(doc, "B. Tables du store jamais écrites (Part III.C)")
gap_table(doc, GH, [
    ["surface_parameters (params SVI par tranche)", "Step 9 + Part V.3",
     "calculés calibration.py:127 ; jamais écrits", "Manquant"],
    ["pricing_results (prix + Greeks par contrat)", "Part III.C",
     "0 store.write(\"pricing_results\")", "Manquant"],
    ["positions (source-of-record versionnée)", "Part III.C",
     "0 store.write(\"positions\")", "Manquant"],
    ["risk_aggregates mal nommée (ligne-à-ligne)", "Step 11 (agrégats groupés)",
     "écrit position_risk_to_dataframe — jobs.py:361", "Partiel"],
], GW)

h2(doc, "C. Réconciliations & checks QC manquants")
gap_table(doc, GH, [
    ["Greeks finite-difference vs analytique", "Step 11 + QC « Greek sanity »",
     "bumps pricing.yaml:12 jamais utilisés ; 0 occurrence", "Manquant"],
    ["Réconciliation Greeks broker vs plateforme", "Step 11",
     "broker greeks captés raw_writer.py:209 ; 0 réconcil.", "Manquant"],
    ["QC check « option chain coverage »", "Part IV.D (10 checks min)",
     "absent de src/qc/checks.py", "Manquant"],
    ["QC check « parity residual »", "Part IV.D",
     "absent de src/qc/checks.py", "Manquant"],
    ["aggregate_risk() câblé + agrégats persistés", "Step 11 (line-level AND aggregate)",
     "def aggregation.py:133 jamais appelé ; .sum() greeks.py:121", "Partiel"],
], GW)

h2(doc, "D. Diagnostics calculés mais non persistés (auditabilité)")
gap_table(doc, GH, [
    ["Diagnostics forward par strike (candidats)", "Step 6 (« keep every candidate row »)",
     "drop dans forward_result_to_dict — engine.py:240", "Partiel"],
    ["Points IV rejetés + warnings de fit", "Step 9 (« store rejected points »)",
     "non écrits — calibration.py:288", "Partiel"],
    ["Payloads broker bruts (instrument master)", "Step 2 (« raw payloads as evidence »)",
     "non stockés — contracts.py:225", "Manquant"],
], GW)

h2(doc, "E. Pipeline & données — comportements incomplets")
gap_table(doc, GH, [
    ["Replay → partitions versionnées physiquement", "Step 13",
     "réécrit le même data.parquet — jobs.py:432", "Partiel"],
    ["Lissage structure par terme + courbe de taux", "Step 6 (e)(f)",
     "taux constant 0.053 ; maturités indépendantes", "Partiel"],
    ["min_open_interest appliqué", "Step 7",
     "seuil qc.yaml jamais filtré — live.py:341", "Partiel"],
    ["Rejet outliers sur IV préliminaires (stade quote)", "Step 7 (d)",
     "seulement dans le forward engine", "Mineur"],
    ["Carry-forward du dernier snapshot fiable", "Step 5",
     "builder fait bid_fallback — builder.py:78", "Mineur"],
    ["exchange_ts renseigné", "Step 3 / data dictionary",
     "toujours None — raw_writer.py:40", "Mineur"],
    ["Quarantaine des events malformés", "Step 3",
     "seulement comptés — raw_writer.py:102", "Mineur"],
    ["Politique de rétention + partition underlying/layer", "Step 4 / Part XV",
     "partition par date seule", "Mineur"],
], GW)

h2(doc, "F. Stratégie de tests — 3 couches sur 4 absentes (Part IV.E)")
para(doc, "Seule la couche 1 (56 tests unitaires) existe.", italic=True, size=9.5, space_after=4)
gap_table(doc, GH, [
    ["Tests d'intégration (adaptateur broker mocké)", "Part IV.E §2",
     "aucun mock IBKR / stream enregistré", "Manquant"],
    ["Tests de régression (fixtures gelées)", "Part IV.E §3 + Part XVI",
     "aucune chaîne figée → forward/IV/surface attendus", "Manquant"],
    ["Tests opérationnels (jobs CLI, alerte sur panne)", "Part IV.E §4",
     "aucun", "Manquant"],
    ["Bibliothèque de journées de replay (calme/stress)", "Part XVI",
     "aucune", "Manquant"],
], GW)

h2(doc, "G. Industrialisation / observabilité (Step 15)")
gap_table(doc, GH, [
    ["Scheduler système (cron/systemd/Airflow)", "Step 15 (a)",
     "boucle in-process — run_collector.py:512", "Manquant"],
    ["Correlation IDs (sessions collecteur ↔ jobs)", "Step 15 (b)",
     "session_id et run_id non tissés ensemble", "Manquant"],
    ["Routage d'alertes externe (email/Slack)", "Step 15 (d)",
     "alerts.json local — run_collector.py:393", "Manquant"],
    ["Catalogue de métriques complet", "Part XIV",
     "partiel (statut collecteur)", "Partiel"],
    ["SLOs formalisés", "Part XIV",
     "sévérités S1–S4 OK, mais pas de cibles SLO", "Mineur"],
], GW)

h2(doc, "H. Points mineurs divers")
gap_table(doc, GH, [
    ["dash + dash-bootstrap-components dans requirements.txt", "Step 1 (reproductibilité)",
     "absents ; environment.md:18 prétend l'inverse", "Partiel"],
    ["Health checks clock-sync + entitlement", "Step 1 (f)",
     "absents du bootstrap", "Mineur"],
    ["Lock file (poetry.lock / pip-tools)", "Step 1 (a)",
     "versions épinglées seulement", "Mineur"],
    ["API de pricing vectorisée", "Step 10",
     "scalaire seulement", "Mineur"],
    ["Accélérateur Newton (annoncé)", "Step 8",
     "docstring solver.py:5 ; non codé", "Mineur"],
    ["Coordonnée delta dans iv_points", "Step 8",
     "absente — solver.py:285", "Mineur"],
    ["README par module + diagramme d'architecture", "Step 16",
     "absents de docs/", "Mineur"],
], GW)

doc.add_page_break()

# ----------------------------------------------------------------------------
# 5. BACKLOG PRIORISÉ
# ----------------------------------------------------------------------------
h1(doc, "6. Backlog priorisé pour fermer l'écart")

h2(doc, "P0 — fort signal, faible coût (à exiger en premier)")
for txt in [
    "Persister surface_parameters (params SVI) — données déjà calculées, table manquante.",
    "Persister les diagnostics par candidat du forward (nouvelle table forward_candidates).",
    "Câbler aggregate_risk() dans le pipeline et écrire de vrais agrégats groupés.",
    "Ajouter dash et dash-bootstrap-components à requirements.txt (corrige l'install front).",
]:
    bullet(doc, txt)

h2(doc, "P1 — checks & réconciliations exigés explicitement")
for txt in [
    "Check QC finite-difference vs analytique (« Greek sanity »).",
    "Checks QC « option chain coverage » et « parity residual ».",
    "Réconciliation Greeks broker vs plateforme (données broker déjà captées).",
    "Eq.22 : fonction interpolate_across_maturities + grille (k, T) continue.",
    "Replay → partitions versionnées physiquement (ne pas écraser data.parquet).",
]:
    bullet(doc, txt)

h2(doc, "P2 — industrialisation & complétude")
for txt in [
    "Tables pricing_results et positions ; politique de rétention.",
    "Couches de tests 2–4 (intégration mockée, régression figée, opérationnels) + journées de replay.",
    "Eq.23 (variance panier), scheduler système, correlation IDs, routage d'alertes externe.",
    "README par module, diagramme d'architecture, schemas.md champ-par-champ.",
]:
    bullet(doc, txt)

para(doc, "")
para(doc, "Conclusion. Le projet est quantitativement sérieux et bien architecturé ; l'écart "
          "résiduel est essentiellement de la PERSISTANCE (3 tables + diagnostics), des "
          "RÉCONCILIATIONS (broker, finite-difference) et de l'INDUSTRIALISATION (tests, "
          "scheduler, observabilité). Les quatre items P0 sont à la fois peu coûteux et "
          "directement vérifiables par un correcteur.", size=10.5, space_after=6)

# ----------------------------------------------------------------------------
# Correctif schéma : le <w:zoom> par défaut de python-docx n'a pas l'attribut
# percent (requis par le schéma). On le fixe à 100 pour une validation propre.
# ----------------------------------------------------------------------------
try:
    _settings = doc.settings.element
    _zoom = _settings.find(qn("w:zoom"))
    if _zoom is None:
        _zoom = OxmlElement("w:zoom"); _settings.append(_zoom)
    if _zoom.get(qn("w:percent")) is None:
        _zoom.set(qn("w:percent"), "100")
except Exception as _exc:
    print(f"zoom fix skipped: {_exc}")

# ----------------------------------------------------------------------------
# Sauvegarde
# ----------------------------------------------------------------------------
OUT = Path(__file__).parent.parent / "docs" / "Audit_conformite_roadmap_v4_detaille.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"OK -> {OUT}")
print(f"Paragraphes: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
