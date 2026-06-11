# -*- coding: utf-8 -*-
"""
Synthèse de l'outil pour le professeur — document de lecture rapide (~10 min) qui
présente ce que contient l'infrastructure, ses subtilités et ses preuves chiffrées.

Génère : Synthese_outil_vol_infrastructure.docx (racine du projet)
Usage   : python scripts/generate_synthese.py
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).parent.parent / "Synthese_outil_vol_infrastructure.docx"

BLUE, GREY, GREEN, AMBER = "1a3a5c", "6a737d", "3f8950", "b58900"


def _shade(cell, color):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(el)


def para(doc, text, bold=False, italic=False, size=10, color=None, space_after=6,
         align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def bullet(doc, text, size=9.5, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
    r = p.add_run(text)
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string(BLUE)
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string("2c5577")
    return h


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = htxt
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        _shade(c, "dbe7f0")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def main():
    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.8)

    # ── Page de garde ──────────────────────────────────────────────────
    para(doc, "", space_after=60)
    para(doc, "Infrastructure de risque de volatilité", bold=True, size=24, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(doc, "EURO STOXX 50 — indice et 50 composantes", size=15, color="2c5577",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "Synthèse de l'outil — lecture ~10 minutes", italic=True, size=12,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    para(doc, "Projet M1 Trading Algorithmique · 11 juin 2026", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    table(doc, ["", "", "", ""],
          [("16/16 étapes de la roadmap\n(audit V6 : 14 Conforme, 2 Conforme−)",
            "25/25 équations\nimplémentées et testées",
            "50/51 sous-jacents collectés\n95,9 % de quotes exploitables",
            "110 tests automatisés\n6 audits de conformité")],
          widths=[1.7, 1.7, 1.7, 1.7])
    doc.add_page_break()

    # ── 1. L'outil en une page ─────────────────────────────────────────
    h1(doc, "1. L'outil en une page")
    para(doc,
         "Une infrastructure institutionnelle de mesure du risque de volatilité sur "
         "l'EURO STOXX 50 : chaque cycle (~30 min), elle collecte ~2 700 options EUREX "
         "sur l'indice et ses 50 composantes via l'API Web d'Interactive Brokers, "
         "reconstruit les forwards par parité put-call, inverse les volatilités "
         "implicites (Black-76 pour l'indice européen, arbre CRR pour les actions "
         "américaines), calibre des surfaces SVI sans arbitrage, produit les greeks "
         "bruts et monétisés en euros, exécute 10 contrôles qualité nommés, et persiste "
         "le tout en Parquet versionné avec une traçabilité complète.", size=10.5)
    para(doc, "Trois commandes suffisent :", bold=True, size=10)
    para(doc, "    1. Gateway IBKR (login navigateur, 1×/session) — https://localhost:5000\n"
              "    2. python run_collector.py        (le collecteur possède la session)\n"
              "    3. python app.py                  (le dashboard lit le store) → localhost:8050",
         size=9.5)
    para(doc,
         "Principe fondateur (process isolation) : le dashboard ne parle JAMAIS au "
         "broker — il lit un store alimenté par le collecteur. Une panne d'UI n'affecte "
         "pas la collecte ; une coupure broker n'éteint pas le dashboard. Les données "
         "brutes sont immuables et toutes les analytics sont recalculables depuis "
         "celles-ci (replay par le même code).", size=10.5)

    # ── 2. Les subtilités qui comptent ─────────────────────────────────
    h1(doc, "2. Les subtilités à connaître (ce qui ne se voit pas au premier regard)")
    bullet(doc, "l'univers est figé en config avec gestion des pièges réels : Sanofi se "
                "résout par le ticker IBKR « SAN1 » (« SAN » seul renvoie l'ADR Santander "
                "NYSE), Nordea par « NDA FI », et trois composantes ont leurs options hors "
                "EUREX (BBVA/IBE sur MEFF, Argenx sur BELFOX) — champ option_exchange par "
                "valeur.", bold_prefix="Résolution des contrats — ")
    bullet(doc, "sélection par DELTA (ATM ± 10Δ/30Δ par aile, call et put), pas par "
                "comptage de strikes ; 7 maturités cibles (1m → 24m) mappées sur les "
                "échéances listées les plus proches, l'écart étant stocké. Ce qui n'existe "
                "pas est flagué « indisponible », jamais inventé.",
           bold_prefix="Grille de collecte — ")
    bullet(doc, "le forward de chaque maturité est extrait de la parité put-call (un "
                "candidat par strike, rejet des outliers par z-score MAD, score de "
                "confiance), et c'est CE forward qui alimente l'inversion d'IV, les greeks "
                "et le pricing — cohérence interne vérifiée par une table de round-trip "
                "(re-pricing à l'IV résolue : erreur médiane 4,5×10⁻⁵ sur l'indice).",
           bold_prefix="Cohérence forward→IV→prix — ")
    bullet(doc, "les greeks sont RECALCULÉS par la plateforme depuis l'IV résolue (jamais "
                "ceux du broker en source de vérité), puis réconciliés deux fois : contre "
                "des différences finies (écart delta médian 0,001) ET contre les greeks "
                "publiés par IBKR (50 pass / 1 skip, écart médian 0,007). Les versions "
                "monétisées € utilisent le multiplicateur réel (10 indice, 100 actions).",
           bold_prefix="Greeks doublement réconciliés — ")
    bullet(doc, "10 checks nommés à seuils versionnés (convergence IV, santé des quotes, "
                "couverture de chaîne, résidu de parité — tolérance ×3 pour les "
                "américaines —, stabilité du forward, bornes du carry, fit de surface, "
                "no-arbitrage calendaire et papillon, réconciliations), baselines "
                "historiques + détection d'anomalies (z-score MAD), triage, et escalade "
                "S1-S4 codée (niveau, responsable, SLA, échéance sur chaque alerte).",
           bold_prefix="Validation comme un produit — ")
    bullet(doc, "chaque écriture conserve une copie versionnée par run "
                "(versions/<run_id>.parquet) ; lineage code_version + config_hash + "
                "run_id sur chaque ligne ; les réponses brutes du broker sont archivées "
                "comme évidence (~9 Mo/jour de JSONL).",
           bold_prefix="Traçabilité totale — ")
    bullet(doc, "limites réelles d'IBKR apprises et gérées en conditions réelles : "
                "~100 lignes de market data simultanées (purge des souscriptions avant "
                "chaque symbole), warm-up du flux différé 3× plus lent à l'ouverture "
                "(adaptatif), 429 sur les endpoints secdef (throttle + retry + caches).",
           bold_prefix="Robustesse opérationnelle — ")

    # ── 3. Le dashboard ────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "3. Le dashboard (15 pages, thème clair desk)")
    table(doc, ["Page", "À voir"],
          [("Market Monitor (accueil)", "Cross-section des 51 : matrice de vol ATM 51×7 cliquable, "
            "niveau vs pente de term structure, moniteur de dispersion, table de l'univers"),
           ("Forward & Carry", "Courbe F(T) + confiance + carry q(T) en % et en € + candidats rejetés"),
           ("Volatilité Implicite", "Smiles superposés toutes maturités, term structure ATM, "
            "skew RR30/BF30, diagnostics du solveur"),
           ("Surface de Vol", "Nappe SVI 3D + heatmap, term structure des paramètres (a,b,ρ,m,σ), "
            "vérification visuelle du no-arbitrage calendaire"),
           ("Greeks & Risk", "La sortie exigée par la spec : grille greeks bruts ET € par "
            "maturité×strike×call/put, lecture desk (+1 % spot), simulateur de choc interactif (Eq.19)"),
           ("QC & Validation", "Jauge, checks individuels, réconciliations diff-finies et broker, triage"),
           ("Pricing / Scénarios / Market Data / Collecteur", "Pricer interactif Black-76 vs CRR · "
            "stress par repricing complet · chaîne brute · monitoring temps réel (heartbeat)")],
          widths=[1.7, 5.2])

    # ── 4. Les bonus quant ─────────────────────────────────────────────
    h1(doc, "4. Au-delà de la roadmap — quatre extensions quantitatives")
    para(doc, "Quatre modules « desk » qui RÉUTILISENT l'infrastructure (spot du store, "
              "σ de la surface SVI, forwards de parité, vega € de la grille) :", size=10.5)
    h2(doc, "4.1 Monte Carlo — option asiatique (mesure risque-neutre)")
    para(doc, "Payoff path-dependent → territoire de MC : GBM sous ℚ avec drift calé point "
              "par point sur NOTRE courbe forward (les forwards de marché sont superposés "
              "aux trajectoires simulées — la moyenne MC passe dessus), quasi-Monte Carlo "
              "(Sobol), variate de contrôle géométrique (Kemna-Vorst exact, variance ÷50), "
              "greeks à nombres aléatoires communs, strike ladder repricé sur les mêmes "
              "chemins. Et un simulateur de DELTA-HEDGE : vente au prix Black puis "
              "couverture discrète le long des chemins — P&L centré sur zéro qui se "
              "resserre avec la fréquence, coûts de transaction en bps, et le pari "
              "vol réalisée vs implicite (E[P&L] ≈ vega·(σimpl − σréal)).", size=10)
    h2(doc, "4.2 Variance swap & mini-VSTOXX (réplication model-free)")
    para(doc, "Le strike de variance répliqué par le log-contrat (méthodologie VIX) sur un "
              "strip de 400 options densifié par la surface SVI ; indice 30 jours interpolé "
              "en variance totale. Résultat : mini-VSTOXX maison ≈ 19,5 — directement "
              "comparable au VSTOXX officiel (champ de saisie pour mesurer l'écart). La "
              "prime de convexité K_var − σ_ATM (le « prix du skew ») est affichée par "
              "maturité, ainsi que le strip EXÉCUTABLE au bid/ask réel.", size=10)
    h2(doc, "4.3 Trade de dispersion (sur la corrélation implicite)")
    para(doc, "La corrélation implicite ρ̄ extraite de l'identité de variance (indice vs "
              "50 composantes, ≈ 0,25-0,35 selon le tenor) devient un trade : short vol "
              "indice / long vol composantes vega-weighted, panier en contrats réels, "
              "P&L vs corrélation réalisée (breakeven au ρ̄ d'entrée), risque net du "
              "package (gamma/theta €) et coût d'entrée au demi-spread des 46 jambes. "
              "Les niveaux ρ̄ et l'indice de variance sont historisés à chaque cycle.", size=10)
    h2(doc, "4.4 Analyseur de risques de stratégies")
    para(doc, "Réplique de l'outil de référence du cours (Stratégies / Mensuel / "
              "Journalier / Intervalle), VALIDÉE AU CENTIÈME contre l'original par 8 tests "
              "(p mensuel rouge 44,26 %, lois binomiale et des séries consécutives par "
              "récurrence dynamique, IC du Sharpe [−0,464 ; 3,464] pour Ŝ=1,5 sur 252 j). "
              "Ajouts : vérification de cohérence (semi-définie-positivité) de la matrice "
              "de corrélation saisie, et courbe de l'IC en fonction de T montrant le "
              "track record nécessaire à la significativité (~1,7 an pour un Sharpe 1,5).",
         size=10)

    # ── 5. La rigueur ──────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "5. La rigueur d'ingénierie (preuves)")
    bullet(doc, "110 tests automatisés : parité put-call au centime, round-trip prix↔IV, "
                "no-arbitrage, martingale du MC, convergence vers les formes fermées, "
                "réplication des sorties de l'outil de référence, callbacks du dashboard.",
           bold_prefix="Tests — ")
    bullet(doc, "six audits de conformité successifs (V2 → V6, Word régénérables par "
                "script), dont un audit INDÉPENDANT sévère (V4 : 9 Conforme / 7 Partiel) "
                "qui a piloté la fermeture de tous les écarts. V6 : 14 Conforme / "
                "2 Conforme− — les 2 réserves restantes sont externes au code (login "
                "navigateur imposé par IBKR retail ; démonstration de handover par un tiers).",
           bold_prefix="Audits — ")
    bullet(doc, "le 11/06 en séance, des lots entiers de cotations sont revenus vides : "
                "diagnostic par hypothèses éliminées avec preuves (session concurrente : "
                "non ; marché fermé : non ; warm-up : partiel) jusqu'à la cause racine "
                "(saturation des souscriptions market data), corrigé et validé en 1h45 — "
                "précisément le critère d'acceptance de l'étape 14 (« un opérateur "
                "identifie les défaillances en minutes »), exercé en réel. Chronologie "
                "complète dans l'audit V6.", bold_prefix="Incident vécu et résolu — ")
    bullet(doc, "runbooks opérationnels (start-of-day, incident market data, ajout de "
                "sous-jacent), checklist de handover automatisée (6/6), scheduler système, "
                "limitations connues DOCUMENTÉES (données différées 15 min, Nordea sans "
                "entitlement Nasdaq Nordic, grille réduite assumée et réversible par "
                "config, poids égaux pour la dispersion).", bold_prefix="Exploitation — ")

    para(doc, "", space_after=2)
    para(doc, "Pour aller plus loin : README.md (vue d'ensemble), methodology.md "
              "(équations 1-25 + méthodes bonus), docs/specification_eurostoxx.md (la "
              "spécification), audit_conformite_roadmap/Audit_conformite_roadmap_v6.docx "
              "(conformité détaillée). Dépôt : github.com/CLDprog/"
              "modelisation_risque_strategie_trading.", italic=True, size=9, color=GREY)
    para(doc, "Document généré par scripts/generate_synthese.py — régénérable.",
         italic=True, size=8, color=GREY)

    doc.save(OUT)
    print(f"Synthese ecrite : {OUT}")


if __name__ == "__main__":
    main()
