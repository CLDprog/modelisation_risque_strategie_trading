from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("Audit_roadmap_vs_projet_architecture_risque.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"
CRITICAL_FILL = "FCE4E4"
HIGH_FILL = "FDEBD0"
MEDIUM_FILL = "FFF4CC"
LOW_FILL = "EAF4E8"


@dataclass
class StepAudit:
    step: str
    title: str
    severity: str
    status: str
    roadmap: str
    current_state: str
    gaps: str
    actions: str


STEP_AUDITS = [
    StepAudit(
        step="1",
        title="Accès, environnements et sécurité",
        severity="Moyen",
        status="Partiel",
        roadmap=(
            "La roadmap exige un environnement reproductible, un bootstrap simple, des secrets hors dépôt, "
            "une configuration externalisée et un smoke test exécutable sans intervention manuelle."
        ),
        current_state=(
            "Le projet possède un bootstrap IBKR, des YAML de configuration et un exemple de `.env`. "
            "Le smoke test de connectivité existe déjà dans `scripts/bootstrap.py`."
        ),
        gaps=(
            "Il manque surtout la couche d'exploitation demandée par la roadmap: vraie documentation d'environnement, "
            "procédure standardisée de provisioning, séparation claire dev/prod-like, et politique explicite de gestion des secrets."
        ),
        actions=(
            "Ajouter une documentation d'installation opérationnelle, figer les variables d'environnement attendues, "
            "définir un mode d'exécution scheduler-friendly et formaliser les prérequis TWS/Gateway."
        ),
    ),
    StepAudit(
        step="2",
        title="Instrument master et discovery",
        severity="Moyen",
        status="Partiel",
        roadmap=(
            "La roadmap veut un instrument master canonique, versionné, reproductible, enrichi par discovery API, "
            "avec persistance de la représentation normalisée et de la réponse broker brute."
        ),
        current_state=(
            "Le module `src/universe/contracts.py` fournit des dataclasses propres, des helpers utiles et une fonction "
            "de discovery IBKR cohérente avec l'architecture cible."
        ),
        gaps=(
            "La persistance des payloads broker bruts n'est pas en place. La version du discovery par date/configuration "
            "n'est pas matérialisée comme couche auditable complète. Le front n'utilise pas réellement cet instrument master."
        ),
        actions=(
            "Persister la réponse brute `reqSecDefOptParams` et la vue canonique, ajouter un identifiant de version de l'univers, "
            "et faire consommer les pages front par le master découvert au lieu d'une chaîne mock."
        ),
    ),
    StepAudit(
        step="3",
        title="Ingestion market data",
        severity="Critique",
        status="Faible",
        roadmap=(
            "La roadmap demande une couche de collecte loss-aware pour sous-jacents et options, append-only, tolérante aux déconnexions, "
            "avec timestamps normalisés, heartbeat, reconnect, métriques de session et possibilité de replay."
        ),
        current_state=(
            "Le projet a bien un `RawMarketEvent`, un `RawEventWriter` et un `LiveCollector`. "
            "La philosophie 'pas d'analytics dans les callbacks' est respectée."
        ),
        gaps=(
            "Le collector actuel ne souscrit en pratique que les underlyings. Les quotes options live, les entitlements failures, "
            "les pacing events et les résumés de couverture par chaîne ne sont pas implémentés au niveau attendu par la roadmap."
        ),
        actions=(
            "Étendre le collector aux options, persister les erreurs de collecte comme événements structurés, "
            "produire un résumé journalier de couverture et garantir le replay d'une journée complète depuis le raw store."
        ),
    ),
    StepAudit(
        step="4",
        title="Stockage persistant et data model",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap exige une séparation claire entre raw immutable et analytics curés, des schémas stables pour chaque couche, "
            "une compatibilité replay/live, une lineage explicite et des règles d'évolution de schéma."
        ),
        current_state=(
            "Le projet dispose d'un `ParquetStore`, d'un `MetadataStore` SQLite, de partitions par date et d'une base raisonnable "
            "pour les couches raw, analytics et manifests."
        ),
        gaps=(
            "Tous les objets dérivés annoncés par la roadmap ne sont pas encore matérialisés comme tables stables. "
            "La provenance `code_version`, `config_hash`, source records et version d'analytics n'est pas propagée de manière uniforme."
        ),
        actions=(
            "Compléter les schémas manquants, enrichir systématiquement chaque partition dérivée avec provenance et versionnement, "
            "et ajouter une convention de recompute/replacement claire pour les partitions dérivées."
        ),
    ),
    StepAudit(
        step="5",
        title="Spot builder et snapshots de marché",
        severity="Critique",
        status="Partiel",
        roadmap=(
            "La roadmap veut des snapshots cohérents, time-aligned, reproductibles à partir du raw, avec spot de référence, "
            "flags de stale/fallback et métriques de complétude."
        ),
        current_state=(
            "Le module `src/snapshots/builder.py` est bien orienté fonctions pures. "
            "Les notions de `reference_spot`, `reference_type`, `quote_age_seconds` et `is_usable` sont déjà présentes."
        ),
        gaps=(
            "Le pipeline de production de `market_state_snapshots` n'est pas réellement branché end-to-end. "
            "Les pages Dash consomment surtout des DataFrames mock recalculés à la volée, pas les snapshots persistés."
        ),
        actions=(
            "Ajouter un vrai job de construction des snapshots depuis le raw, persister les partitions correspondantes, "
            "et basculer les pages `Market Data` et `Implied Vol` sur cette source plutôt que sur `src.data.mock`."
        ),
    ),
    StepAudit(
        step="6",
        title="Forward et carry",
        severity="Moyen",
        status="Partiel",
        roadmap=(
            "La roadmap exige un forward robuste par maturité, construit depuis plusieurs strikes proches de l'ATM, "
            "avec pondération de liquidité, rejet d'outliers, diagnostics complets et courbe de carry implicite."
        ),
        current_state=(
            "Le moteur `src/forwards/engine.py` est globalement aligné: parité put-call, poids de liquidité, z-score robuste MAD, "
            "score de confiance et carry implicite."
        ),
        gaps=(
            "Le moteur est bon en isolation mais pas encore alimenté par une vraie chaîne live filtrée. "
            "Le front `forward.py` lit une courbe mock ou semi-synthétique au lieu de refléter les outputs du moteur analytique."
        ),
        actions=(
            "Raccorder le moteur aux snapshots persistés, stocker la courbe forward et les diagnostics détaillés par maturité, "
            "puis faire consommer cette sortie au front et au pipeline EOD."
        ),
    ),
    StepAudit(
        step="7",
        title="Quote normalization et QC des quotes",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap demande une vraie politique de sélection des quotes avec reason codes auditables, "
            "filtered chain, full chain conservée, checks nommés et stabilité de comportement sous une version de seuils donnée."
        ),
        current_state=(
            "Le projet possède déjà des critères utiles: bid positif, marché croisé/bloqué, âge max, `is_usable`, `reject_reason`, "
            "et un module QC séparé."
        ),
        gaps=(
            "La normalisation des quotes n'est pas encore une couche autonome pleinement auditée comme l'exige la roadmap. "
            "Les motifs de rejet existent mais ne sont pas encore déclinés en bibliothèque exhaustive, versionnée, branchée à un vrai flux live options."
        ),
        actions=(
            "Isoler la filtered chain comme dataset dédié, conserver la full chain brute de snapshot, "
            "enrichir les checks avec monotonicité, prix impossibles vs intrinsèque, OI/volume et outliers pré-IV."
        ),
    ),
    StepAudit(
        step="8",
        title="Solveur d'IV",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap veut un solveur robuste, bracketed, diagnostic, avec traitement propre des non-convergences, "
            "support batch et convention claire pour les cas américains."
        ),
        current_state=(
            "Le solveur `src/iv/solver.py` est bien structuré et proprement typé: bornes, résidu, itérations, diagnostics et batch wrapper."
        ),
        gaps=(
            "Le support des options américaines n'est pas réellement traité côté inversion. "
            "Le solveur n'est pas encore branché à un flux de quotes filtrées industrielles, et le front affiche des IV mock déjà prêtes."
        ),
        actions=(
            "Définir la convention d'IV pour produits américains, brancher le solveur sur la filtered chain persistée, "
            "et stocker la table d'IV comme vraie sortie canonique consommée par les pages et le risk engine."
        ),
    ),
    StepAudit(
        step="9",
        title="Surface engine",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap exige une séparation stricte entre points IV résolus et surface fitted, "
            "stockage des paramètres et de la grille, diagnostics de fit, alertes de sparsité et contrôle des arbitrages statiques."
        ),
        current_state=(
            "Le module `src/surfaces/calibration.py` couvre SVI, fallback spline, RMSE et monotonicité calendaire, ce qui constitue une base sérieuse."
        ),
        gaps=(
            "Le front `pages/surface.py` n'utilise pas réellement le moteur de surface calibré. "
            "Les contrôles d'arbitrage restent limités. Les opérateurs ne comparent pas encore raw points vs fitted surface dans l'interface."
        ),
        actions=(
            "Faire converger la page surface vers les `surface_results`, persister paramètres et grille, "
            "ajouter des diagnostics visuels raw-vs-fit et enrichir les checks de static arbitrage au-delà de la seule monotonicité calendaire."
        ),
    ),
    StepAudit(
        step="10",
        title="Pricing engine",
        severity="Moyen",
        status="Partiel",
        roadmap=(
            "La roadmap veut un service de pricing central, réutilisable, cohérent avec l'inversion, avec API stable, "
            "Greeks, benchs, conventions documentées et éventuelle interface vectorisée."
        ),
        current_state=(
            "Les pricers européen et américain sont propres, testés, documentés et séparés. "
            "Le projet a déjà une base de tests convaincante sur ces briques."
        ),
        gaps=(
            "Il manque surtout la maturation industrielle: interface vectorisée claire, benchmarks de performance, "
            "et intégration systématique du pricing engine comme unique source de vérité de valorisation dans tout le projet."
        ),
        actions=(
            "Formaliser une API de batch pricing, ajouter des tests de performance et documenter précisément "
            "les hypothèses par famille de produit."
        ),
    ),
    StepAudit(
        step="11",
        title="Greeks et risk analytics",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap demande une couche de risque canonique à partir de vraies positions, avec outputs ligne et agrégés, "
            "conventions stables, et réconciliation possible avec des Greeks broker si disponibles."
        ),
        current_state=(
            "Les structures `Position` et `PositionRisk` existent, avec agrégation et calculs monétisés cohérents."
        ),
        gaps=(
            "Le front `greeks.py` reste branché sur un portefeuille mock. "
            "Il n'y a pas encore d'intégration source-of-record des positions ni de réconciliation broker Greeks."
        ),
        actions=(
            "Brancher une vraie source de positions, persister les risk snapshots, ajouter un rapport de réconciliation "
            "et faire consommer cette sortie au dashboard de Greeks."
        ),
    ),
    StepAudit(
        step="12",
        title="Scénarios et diagnostics margin-style",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap impose une grille de scénarios versionnée, repricing complet comme vérité, approximation locale comme aide, "
            "et une sortie reproductible par snapshot, positions et version de scénario."
        ),
        current_state=(
            "Le moteur de scénarios est bien cadré mathématiquement et la grille YAML existe déjà."
        ),
        gaps=(
            "Le pipeline EOD appelle `run_all_scenarios` avec une liste de positions vide. "
            "La page `scenarios.py` lit des résultats mock et non la vraie sortie du moteur branché à des positions réelles."
        ),
        actions=(
            "Injecter les vraies positions dans l'orchestration, stocker les résultats de scénarios comme table canonique, "
            "et faire consommer cette table au front au lieu d'un générateur mock."
        ),
    ),
    StepAudit(
        step="13",
        title="Historical reconstruction et replay",
        severity="Critique",
        status="Absent",
        roadmap=(
            "La roadmap veut un replay complet d'une journée ou d'un mois via le même code path que le live, "
            "avec gestion explicite des partitions manquantes et archivage versionné des restatements."
        ),
        current_state=(
            "Le design du projet favorise le replay théorique, car plusieurs briques sont pures et déterministes."
        ),
        gaps=(
            "Le workflow de replay n'est pas réellement implémenté comme produit. "
            "Il manque les scripts, les jobs batch et les rapports comparant replay et live sur périodes chevauchantes."
        ),
        actions=(
            "Construire un job de replay de bout en bout, matérialiser les partitions restatées avec version analytique, "
            "et ajouter un rapport de comparaison replay-vs-live."
        ),
    ),
    StepAudit(
        step="14",
        title="Validation framework et anomaly detection",
        severity="Critique",
        status="Partiel",
        roadmap=(
            "La roadmap veut un produit de validation: checks nommés, pass/warn/fail, contextes détaillés, "
            "table de triage, baseline historique, anomalies et escalade opérateur."
        ),
        current_state=(
            "Le projet a déjà un bon squelette de checks nommés et un format `QcResult` clair."
        ),
        gaps=(
            "Il manque encore la couche de tendance/anomalie, la table de triage structurée, la vue opérateur historique, "
            "et les mécanismes d'escalade. La page QC actuelle reste mockée."
        ),
        actions=(
            "Brancher les résultats QC réels au front, stocker les échecs par instrument/maturité, "
            "ajouter des baselines glissantes et construire un dashboard de suivi QC historique."
        ),
    ),
    StepAudit(
        step="15",
        title="Orchestration, logging et observabilité",
        severity="Critique",
        status="Partiel",
        roadmap=(
            "La roadmap exige des jobs planifiés, des retry policies, des correlation IDs, des alertes, "
            "des métriques d'exploitation et des redémarrages idempotents."
        ),
        current_state=(
            "Le projet a un `run_eod_pipeline`, du logging `loguru`, des manifests et un bootstrap. "
            "La séparation des modules facilite l'orchestration future."
        ),
        gaps=(
            "Il manque les vrais wrappers scheduler, les alertes, la supervision métier, les corrélations entre sessions de collecte "
            "et runs analytiques, ainsi qu'une gestion robuste du backlog et des restart procedures."
        ),
        actions=(
            "Définir les jobs opérés, enrichir les manifests et logs avec IDs corrélés, exposer des métriques "
            "sur la collecte et les analytics, puis brancher une couche d'alerte."
        ),
    ),
    StepAudit(
        step="16",
        title="Production handover, documentation et SOP",
        severity="Élevé",
        status="Partiel",
        roadmap=(
            "La roadmap exige une vraie documentation d'exploitation: interfaces gelées, runbooks, release checklist, "
            "ownership, limitations connues et handover démontrable par un autre opérateur."
        ),
        current_state=(
            "Le dépôt a un `README`, un `AGENTS.md`, un `CLAUDE.md` et des commentaires utiles dans les modules."
        ),
        gaps=(
            "Il manque encore la documentation d'exploitation demandée par la roadmap: runbooks ciblés, SOP incident/replay/QC, "
            "checklist de release, matrice de support et démonstration de handover."
        ),
        actions=(
            "Créer un vrai dossier `docs/` orienté opérations, documenter les procédures récurrentes, "
            "et formaliser un guide de prise en main pour un opérateur ou développeur junior."
        ),
    ),
]


COHERENCE_GAPS = [
    (
        "Le front Dash ne consomme pas encore la chaîne analytique industrielle.",
        "Plusieurs pages lisent des générateurs mock ou des reconstructions ad hoc au lieu des partitions persistées issues du pipeline backend."
    ),
    (
        "Le mode live n'est pas vraiment live pour la chaîne d'options.",
        "Le spot vient d'IBKR, mais la chaîne d'options est reconstruite synthétiquement autour du spot live; ce n'est pas la même chose qu'une collecte d'observations options réelles."
    ),
    (
        "Le pipeline EOD n'est pas complètement fermé.",
        "Il lit des snapshots supposés exister, mais la chaîne de production de ces snapshots n'est pas encore industrialisée de bout en bout dans le projet."
    ),
    (
        "Les sorties risk, scenarios et QC ne sont pas encore les sorties canoniques du système.",
        "Le projet possède les moteurs métier, mais les vues utilisateur restent surtout des démonstrateurs branchés sur du mock."
    ),
    (
        "La provenance et l'auditabilité ne sont pas encore homogènes.",
        "La roadmap exige la traçabilité complète des objets dérivés; aujourd'hui le manifest existe, mais pas encore la discipline uniforme de lineage pour toutes les tables."
    ),
    (
        "La dimension exploitation/ops est en retard par rapport au code quantitatif.",
        "Les briques analytiques sont plus avancées que l'observabilité, les runbooks, les alertes et les procédures de support."
    ),
]


PRIORITY_ACTIONS = [
    (
        "Priorité 1 — Fermer la chaîne live réelle",
        "Collecter et persister les quotes options live, construire les snapshots de marché depuis le raw, "
        "et supprimer la dépendance fonctionnelle du front aux générateurs mock pour les pages coeur."
    ),
    (
        "Priorité 2 — Faire converger le front et le backend",
        "Brancher `market_data`, `implied_vol`, `surface`, `greeks`, `scenarios` et `qc` sur les sorties persistées du pipeline "
        "au lieu d'utiliser des reconstructions ou datasets mock."
    ),
    (
        "Priorité 3 — Industrialiser le pipeline analytique",
        "Ajouter les jobs manquants entre raw events et snapshots, brancher les vraies positions, exécuter réellement les scénarios, "
        "et persister chaque couche comme objet canonique versionné."
    ),
    (
        "Priorité 4 — Renforcer validation et observabilité",
        "Construire la table de triage QC, les baselines historiques, les alertes, les manifests enrichis, "
        "et les métriques d'exploitation de collecte, solver, surface et scénarios."
    ),
    (
        "Priorité 5 — Mettre en place replay et handover",
        "Créer un workflow de replay/backfill same-code-path, puis documenter l'exploitation dans des runbooks utilisables "
        "par une autre personne sans dépendre du développeur initial."
    ),
]


FILE_APPENDIX = [
    "`src/connectivity/session.py` : bonne base de session IBKR et de reconnect.",
    "`src/universe/contracts.py` : instrument master propre mais discovery brute à mieux persister.",
    "`src/collectors/raw_writer.py` : raw event layer cohérente, mais collecte options live incomplète.",
    "`src/snapshots/builder.py` : design pur et reproductible, mais job de production des snapshots à fermer.",
    "`src/forwards/engine.py` : moteur forward crédible et bien orienté diagnostics.",
    "`src/iv/solver.py` : solveur robuste, déjà proche des exigences de la roadmap.",
    "`src/surfaces/calibration.py` : base SVI/spline sérieuse, encore sous-utilisée par le front.",
    "`src/risk/aggregation.py` et `src/risk/scenarios.py` : bonnes briques métier, mais trop peu branchées sur de vraies positions/sorties pipeline.",
    "`src/qc/checks.py` : bon squelette de QC, à enrichir en produit de validation opérable.",
    "`src/orchestration/jobs.py` : squelette EOD utile, mais pipeline encore incomplet.",
    "`src/data/mock.py`, `src/data/live.py`, `src/data/source.py` : zone centrale de la divergence actuelle entre démonstration front et architecture industrielle."
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    dxa = int(width_inches * 1440)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_inches)


def set_table_width(table, width_inches: float = 6.5, indent_inches: float = 0.083) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(int(indent_inches * 1440)))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_run_font(run, name="Calibri", size=11, color=TEXT, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.10


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Audit de cohérence roadmap vs projet")
    set_run_font(header_run, size=9, color=MUTED, bold=False)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Page ")
    set_run_font(footer_run, size=9, color=MUTED)
    add_page_number(footer)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Audit détaillé de cohérence")
    set_run_font(run, size=23, color=TEXT, bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(14)
    run = p2.add_run("Roadmap v4.0 vs implémentation actuelle du projet `architecture_risque`")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("Document", "Audit fonctionnel et technique"),
        ("Périmètre", "Roadmap PDF du 06 avril 2026, codebase actuelle, configuration, pages Dash, backend `src/`, tests"),
        ("Date", date.today().strftime("%d/%m/%Y")),
        ("Objectif", "Identifier précisément les écarts à la roadmap et lister les actions nécessaires pour rendre le projet cohérent avec les exigences demandées"),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, 6.5, 0.083)
    widths = [1.7, 4.8]
    for label, value in metadata:
        row = table.add_row().cells
        for idx, width in enumerate(widths):
            set_cell_width(row[idx], width)
            row[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row[0], HEADER_FILL)
        p_label = row[0].paragraphs[0]
        p_label.paragraph_format.space_after = Pt(2)
        p_label.paragraph_format.space_before = Pt(2)
        p_val = row[1].paragraphs[0]
        p_val.paragraph_format.space_after = Pt(2)
        p_val.paragraph_format.space_before = Pt(2)
        set_run_font(p_label.add_run(label), size=10.5, color=TEXT, bold=True)
        set_run_font(p_val.add_run(value), size=10.5, color=TEXT)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        set_run_font(p.add_run(bold_prefix), size=11, color=TEXT, bold=True)
        set_run_font(p.add_run(text), size=11, color=TEXT)
    else:
        set_run_font(p.add_run(text), size=11, color=TEXT)


def add_section_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    style_doc(doc)
    add_header_footer(doc)


def add_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Synthèse exécutive", level=1)
    add_paragraph(
        doc,
        "Le projet est globalement bien inspiré par la roadmap et reprend correctement ses grandes briques: "
        "connectivité IBKR, instrument master, raw events, snapshots, forward, solveur d'IV, surface, pricing, risk, QC et orchestration."
    )
    add_paragraph(
        doc,
        "Le point central de l'audit est le suivant: l'architecture conceptuelle est déjà proche de la cible, "
        "mais l'exécution bout en bout reste incomplète. Aujourd'hui, la partie la plus avancée du projet est le noyau quantitatif et analytique; "
        "la partie la moins mature est la fermeture industrielle de la chaîne de données, l'exploitation opérateur et la cohérence front/back."
    )
    add_paragraph(
        doc,
        "En pratique, le dépôt ressemble davantage à un socle de démonstration avancée et de prototypage structuré qu'à une plateforme "
        "de volatilité totalement conforme au blueprint 'institutional-grade' de la roadmap."
    )
    add_paragraph(
        doc,
        "La divergence la plus importante vient du fait que plusieurs pages Dash consomment encore des données mock ou semi-synthétiques, "
        "alors que la roadmap exige une chaîne traçable 'raw observations -> market-state snapshots -> analytics -> risk -> QC'."
    )


def add_overall_matrix(doc: Document) -> None:
    doc.add_heading("2. Tableau de synthèse des écarts", level=1)
    add_paragraph(
        doc,
        "Le tableau ci-dessous résume l'état de chaque étape de la roadmap par rapport au projet actuel. "
        "Le statut doit être lu de manière opérationnelle: 'Partiel' signifie que la bonne brique existe mais n'est pas encore fermée end-to-end; "
        "'Faible' signifie que l'intention existe mais que la fonctionnalité cible n'est pas encore réellement livrée; "
        "'Absent' signifie que la roadmap demande un produit ou un workflow qui n'est pas encore en place."
    )

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, 6.5, 0.083)
    widths = [0.45, 1.55, 0.95, 0.85, 2.70]
    headers = ["Étape", "Workstream", "Statut", "Sévérité", "Manque principal"]
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(headers[i]), size=10, color=TEXT, bold=True)

    severity_fill = {
        "Critique": CRITICAL_FILL,
        "Élevé": HIGH_FILL,
        "Moyen": MEDIUM_FILL,
        "Faible": LOW_FILL,
    }
    main_gap = {
        "1": "Ops et sécurité encore trop peu formalisés.",
        "2": "Persistance brute du discovery et versionnement incomplets.",
        "3": "Collecte live options manquante.",
        "4": "Lineage et schémas dérivés pas encore homogènes.",
        "5": "Snapshots non fermés en vrai pipeline consommé par le front.",
        "6": "Moteur bon, mais peu branché au flux réel.",
        "7": "Filtered chain et bibliothèque QC encore incomplètes.",
        "8": "IV branchée au mock; convention américaine incomplète.",
        "9": "Surface moteur sous-utilisée par l'UI.",
        "10": "Industrialisation API/perf encore limitée.",
        "11": "Pas de vraies positions branchées.",
        "12": "Scénarios réels non exécutés dans l'orchestration.",
        "13": "Replay historique absent.",
        "14": "Validation produit/opérateur incomplète.",
        "15": "Observabilité et alerting encore insuffisants.",
        "16": "Runbooks et handover à construire.",
    }
    for item in STEP_AUDITS:
        row = table.add_row().cells
        values = [item.step, item.title, item.status, item.severity, main_gap[item.step]]
        for i, cell in enumerate(row):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 3:
                set_cell_shading(cell, severity_fill.get(item.severity, LOW_FILL))
            p = cell.paragraphs[0]
            if i in (0, 2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(values[i]), size=9.5, color=TEXT, bold=(i == 0))


def add_coherence_gap_section(doc: Document) -> None:
    doc.add_heading("3. Incohérences structurantes entre la roadmap et le projet", level=1)
    add_paragraph(
        doc,
        "Avant de détailler les 16 étapes, il est utile d'isoler les six incohérences systémiques qui expliquent la majorité des écarts observés. "
        "Ces incohérences ne sont pas des détails de code: ce sont les raisons pour lesquelles le projet n'est pas encore complètement cohérent "
        "avec le modèle d'exploitation demandé par la roadmap."
    )
    for idx, (title, detail) in enumerate(COHERENCE_GAPS, start=1):
        doc.add_heading(f"3.{idx} {title}", level=2)
        add_paragraph(doc, detail)


def add_step_details(doc: Document) -> None:
    doc.add_heading("4. Audit détaillé par étape de roadmap", level=1)
    add_paragraph(
        doc,
        "Chaque sous-section reprend explicitement: ce que la roadmap veut et exige, "
        "ce que l'implémentation actuelle fait réellement, ce qui manque encore, et ce qu'il faut faire pour remettre l'étape en cohérence."
    )
    for item in STEP_AUDITS:
        doc.add_heading(f"4.{item.step} Étape {item.step} — {item.title}", level=2)
        add_paragraph(doc, item.roadmap, "Exigence de la roadmap. ")
        add_paragraph(doc, item.current_state, "État actuel du projet. ")
        add_paragraph(doc, item.gaps, "Écart constaté. ")
        add_paragraph(doc, item.actions, "Travail à mener. ")


def add_priority_plan(doc: Document) -> None:
    doc.add_heading("5. Liste priorisée des choses à faire", level=1)
    add_paragraph(
        doc,
        "La liste suivante est ordonnée par impact structurel. L'idée n'est pas d'ajouter des fonctionnalités isolées, "
        "mais de transformer le projet pour qu'il respecte réellement la logique de la roadmap."
    )
    for idx, (title, detail) in enumerate(PRIORITY_ACTIONS, start=1):
        doc.add_heading(f"5.{idx} {title}", level=2)
        add_paragraph(doc, detail)

    doc.add_heading("5.6 Ordre d'exécution recommandé", level=2)
    add_paragraph(
        doc,
        "L'ordre recommandé est le suivant: 1) fermer la collecte et les snapshots réels, 2) brancher les analytics persistées aux pages front, "
        "3) intégrer les vraies positions et les scénarios, 4) construire la validation et l'observabilité, 5) ajouter le replay et la documentation d'exploitation."
    )
    add_paragraph(
        doc,
        "Si cet ordre n'est pas respecté, le projet risque de rester longtemps dans une zone intermédiaire: "
        "beau sur le plan algorithmique, mais insuffisamment traçable et insuffisamment exploitable au regard du cahier des charges."
    )


def add_target_state(doc: Document) -> None:
    doc.add_heading("6. Ce que doit devenir le projet pour être cohérent avec la roadmap", level=1)
    add_paragraph(
        doc,
        "À l'arrivée, le projet doit fonctionner comme une chaîne continue et auditée: "
        "IBKR -> raw events immuables -> snapshots de marché persistés -> quote filtering -> forward -> IV -> surface -> pricing -> risk -> scenarios -> QC -> reporting."
    )
    add_paragraph(
        doc,
        "Le front Dash ne doit plus être un producteur parallèle de données de démonstration. "
        "Il doit devenir une couche de visualisation fidèle des objets canoniques générés par le backend."
    )
    add_paragraph(
        doc,
        "Chaque résultat affiché à l'utilisateur doit être explicable par une partition persistée, "
        "une version de code, une version de configuration et un historique de calcul documenté."
    )
    add_paragraph(
        doc,
        "Autrement dit, le projet sera cohérent avec la roadmap non pas seulement lorsque chaque module existera, "
        "mais lorsque l'ensemble sera traçable, rejouable, exploitable et lisible par une autre personne que le développeur initial."
    )


def add_appendix(doc: Document) -> None:
    doc.add_heading("Annexe — fichiers clés à traiter en priorité", level=1)
    for line in FILE_APPENDIX:
        add_paragraph(doc, line)


def build_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_header_footer(doc)
    add_title_block(doc)
    add_executive_summary(doc)
    add_overall_matrix(doc)
    add_coherence_gap_section(doc)
    add_step_details(doc)
    add_priority_plan(doc)
    add_target_state(doc)
    add_appendix(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path.resolve())
